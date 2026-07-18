import unittest
from io import BytesIO

from docx import Document as DocxDocument
from fastapi import FastAPI
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from routes.assignment_plans import router
from services.auth import get_current_user
from services.document_extractor import (
    DocumentExtractionError,
    UploadedDocument,
    extract_documents,
    safe_filename,
)
from services.request_limits import RequestBodyLimitMiddleware


def uploaded(name: str, data: bytes) -> UploadedDocument:
    return UploadedDocument(name=name, content_type=None, data=data)


def docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("Write a 1,500 word analysis with three academic sources.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Criterion"
    table.cell(0, 1).text = "Weight"
    table.cell(1, 0).text = "Evidence"
    table.cell(1, 1).text = "40%"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def merged_table_docx_bytes() -> bytes:
    document = DocxDocument()
    table = document.add_table(rows=1, cols=3)
    merged = table.cell(0, 0).merge(table.cell(0, 2))
    merged.text = "One merged grading criterion"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_bytes(text: str | None = None) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if text:
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): writer._add_object(font)}
                )
            }
        )
        content = DecodedStreamObject()
        content.set_data(
            f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
        )
        page[NameObject("/Contents")] = writer._add_object(content)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


class DocumentExtractorTests(unittest.TestCase):
    def test_extracts_multiple_text_documents_in_order(self):
        response = extract_documents(
            [
                uploaded("prompt.txt", b"Write the final report."),
                uploaded("rubric.txt", b"Use evidence and cite every source."),
            ]
        )

        self.assertLess(response.text.index("final report"), response.text.index("cite every"))
        self.assertIn("Document 1 (TXT)", response.text)
        self.assertEqual([item.name for item in response.files], ["prompt.txt", "rubric.txt"])

    def test_extracts_pdf_text(self):
        response = extract_documents(
            [uploaded("assignment.pdf", pdf_bytes("Explain the experimental results."))]
        )

        self.assertIn("experimental results", response.text)
        self.assertEqual(response.files[0].pageCount, 1)

    def test_rejects_pdf_without_a_text_layer(self):
        with self.assertRaisesRegex(DocumentExtractionError, "no extractable text"):
            extract_documents([uploaded("scan.pdf", pdf_bytes())])

    def test_extracts_docx_paragraphs_and_tables(self):
        response = extract_documents([uploaded("rubric.docx", docx_bytes())])

        self.assertIn("1,500 word analysis", response.text)
        self.assertIn("Criterion | Weight", response.text)
        self.assertIn("Evidence | 40%", response.text)

    def test_extracts_merged_docx_cell_once(self):
        response = extract_documents(
            [uploaded("merged-rubric.docx", merged_table_docx_bytes())]
        )

        self.assertEqual(response.text.count("One merged grading criterion"), 1)

    def test_accepts_bom_marked_utf16_text(self):
        data = "Include an annotated bibliography.".encode("utf-16")

        response = extract_documents([uploaded("notes.txt", data)])

        self.assertIn("annotated bibliography", response.text)

    def test_rejects_binary_text_file(self):
        with self.assertRaisesRegex(DocumentExtractionError, "binary file"):
            extract_documents([uploaded("notes.txt", b"brief\x00binary")])

    def test_rejects_unsupported_extensions(self):
        with self.assertRaisesRegex(DocumentExtractionError, "not supported"):
            extract_documents([uploaded("legacy.doc", b"content")])

    def test_does_not_truncate_when_character_budget_is_exceeded(self):
        with self.assertRaisesRegex(DocumentExtractionError, "exceeds"):
            extract_documents([uploaded("brief.txt", b"a" * 200)], max_characters=100)

    def test_strips_paths_and_control_characters_from_file_names(self):
        self.assertEqual(safe_filename("C:\\fake\\rubric\x00.txt"), "rubric.txt")

    def test_preserves_supported_extension_when_truncating_file_name(self):
        name = safe_filename(f"{'a' * 150}.pdf")

        self.assertEqual(len(name), 120)
        self.assertTrue(name.endswith(".pdf"))


class DocumentExtractionRouteTests(unittest.TestCase):
    def setUp(self):
        app = FastAPI()
        app.include_router(router, prefix="/assignment-plans")
        app.dependency_overrides[get_current_user] = lambda: {"uid": "test-user"}
        self.client = TestClient(app)

    def test_extracts_repeated_multipart_files(self):
        response = self.client.post(
            "/assignment-plans/extract",
            files=[
                ("files", ("prompt.txt", b"Write a report.", "text/plain")),
                ("files", ("rubric.txt", b"Cite three sources.", "text/plain")),
            ],
            data={"maxCharacters": "50000"},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.json()["files"]), 2)

    def test_rejects_more_than_five_files(self):
        response = self.client.post(
            "/assignment-plans/extract",
            files=[
                ("files", (f"brief-{index}.txt", b"brief", "text/plain"))
                for index in range(6)
            ],
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("five", response.json()["detail"])

    def test_rejects_unsupported_file_type(self):
        response = self.client.post(
            "/assignment-plans/extract",
            files=[("files", ("brief.rtf", b"brief", "application/rtf"))],
        )

        self.assertEqual(response.status_code, 415)

    def test_rejects_request_body_before_multipart_spooling_exceeds_limit(self):
        app = FastAPI()
        app.add_middleware(
            RequestBodyLimitMiddleware,
            path="/assignment-plans/extract",
            max_body_bytes=512,
        )
        app.include_router(router, prefix="/assignment-plans")
        app.dependency_overrides[get_current_user] = lambda: {"uid": "test-user"}

        response = TestClient(app).post(
            "/assignment-plans/extract",
            files=[("files", ("brief.txt", b"a" * 1024, "text/plain"))],
        )

        self.assertEqual(response.status_code, 413)

    def test_authentication_runs_before_request_body_is_parsed(self):
        app = FastAPI()
        app.add_middleware(
            RequestBodyLimitMiddleware,
            path="/assignment-plans/extract",
            max_body_bytes=1,
        )
        app.include_router(router, prefix="/assignment-plans")

        response = TestClient(app).post(
            "/assignment-plans/extract",
            files=[("files", ("brief.txt", b"brief", "text/plain"))],
        )

        self.assertEqual(response.status_code, 401)

    def test_rejects_oversized_non_multipart_body_with_413(self):
        app = FastAPI()
        app.add_middleware(
            RequestBodyLimitMiddleware,
            path="/assignment-plans/extract",
            max_body_bytes=32,
        )
        app.include_router(router, prefix="/assignment-plans")
        app.dependency_overrides[get_current_user] = lambda: {"uid": "test-user"}

        response = TestClient(app).post(
            "/assignment-plans/extract",
            content=b"a" * 64,
            headers={"Content-Type": "application/x-www-form-urlencoded"},
        )

        self.assertEqual(response.status_code, 413)


if __name__ == "__main__":
    unittest.main()
